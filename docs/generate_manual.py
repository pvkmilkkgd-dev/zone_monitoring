"""
Скрипт для генерации инструкции пользователя в формате Word (.docx)
Запуск: python generate_manual.py
Требуется: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_manual():
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Заголовок документа
    title = doc.add_heading('Инструкция пользователя', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Система мониторинга обстановки')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Содержание
    doc.add_heading('Содержание', level=1)
    contents = [
        '1. Общие сведения',
        '2. Роли пользователей',
        '3. Вход в систему',
        '4. Страница обстановки',
        '5. Работа с событиями',
        '6. Работа со слоями',
        '7. Отчёты',
        '8. Администрирование',
    ]
    for item in contents:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Общие сведения
    doc.add_heading('1. Общие сведения', level=1)
    doc.add_paragraph('Система мониторинга обстановки предназначена для:')
    features = [
        'Визуального отображения событий на карте региона',
        'Управления событиями различной важности',
        'Категоризации событий по слоям и подслоям',
        'Формирования отчётов и выгрузки данных',
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 2. Роли пользователей
    doc.add_heading('2. Роли пользователей', level=1)
    doc.add_paragraph('В системе предусмотрены следующие роли:')
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    # Заголовок таблицы
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Роль'
    header_cells[1].text = 'Описание'
    header_cells[2].text = 'Доступные функции'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    roles_data = [
        ('Администратор (admin)', 'Полный доступ ко всем функциям', 'Настройки системы, управление пользователями, зонами, слоями, событиями, журнал аудита'),
        ('Старший редактор (editor_plus)', 'Расширенные права редактирования', 'Редактирование всех событий, слоёв, формирование отчётов'),
        ('Редактор (editor)', 'Базовые права редактирования', 'Создание событий, редактирование только своих событий'),
        ('Наблюдатель (viewer)', 'Только просмотр', 'Просмотр карты обстановки'),
    ]
    
    for i, (role, desc, funcs) in enumerate(roles_data, 1):
        row = table.rows[i].cells
        row[0].text = role
        row[1].text = desc
        row[2].text = funcs
    
    # 3. Вход в систему
    doc.add_heading('3. Вход в систему', level=1)
    steps = [
        'Откройте браузер и перейдите по адресу системы',
        'Введите ваш логин и пароль',
        'Нажмите кнопку «Войти»',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph()
    doc.add_paragraph('После успешного входа вы будете перенаправлены:')
    redirects = [
        'Администратор → Страница настроек (/admin)',
        'Редактор / Старший редактор → Страница событий (/editor/events)',
        'Наблюдатель → Страница обстановки (/situation)',
    ]
    for r in redirects:
        doc.add_paragraph(r, style='List Bullet')
    
    note = doc.add_paragraph()
    note.add_run('Примечание: ').bold = True
    note.add_run('Сессия действует 2 часа. После истечения этого времени потребуется повторный вход.')
    
    # 4. Страница обстановки
    doc.add_heading('4. Страница обстановки (/situation)', level=1)
    doc.add_paragraph('Главная страница с интерактивной картой региона.')
    
    doc.add_heading('Элементы интерфейса:', level=2)
    
    doc.add_paragraph('Верхняя панель:', style='List Bullet')
    upper_panel = [
        'Фильтр по слоям — выбор отображаемых категорий событий',
        'Название управления — отображается по центру',
        'Кнопка «Панель управления» (иконка домика) — переход к редактированию (только для редакторов и администраторов)',
        'Кнопка «Выход» — выход из системы',
    ]
    for item in upper_panel:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_paragraph('Карта:', style='List Bullet')
    map_info = [
        'Районы окрашены в зависимости от важности событий:',
        '  • Зелёный — низкая важность (1-3)',
        '  • Жёлтый — средняя важность (4-6)',
        '  • Красный — высокая важность (7-10)',
        'При клике на район открывается боковая панель с информацией',
    ]
    for item in map_info:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_paragraph('Боковая панель (при выборе района):', style='List Bullet')
    side_panel = [
        'Название района и подразделения',
        'Список событий в данном районе',
        'При клике на событие — подробная информация и комментарии',
    ]
    for item in side_panel:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_paragraph('Нижний левый угол:', style='List Bullet')
    bottom_left = [
        'Легенда важности событий',
        'Имя текущего пользователя',
    ]
    for item in bottom_left:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    # 5. Работа с событиями
    doc.add_heading('5. Работа с событиями', level=1)
    doc.add_paragraph('Доступно для: администратор, старший редактор, редактор').italic = True
    
    doc.add_heading('Страница событий (/editor/events)', level=2)
    
    doc.add_heading('Создание события:', level=3)
    create_steps = [
        'Нажмите «Создать событие»',
        'Заполните поля:',
        '   • Название — краткое описание события (обязательно)',
        '   • Важность — от 1 до 10 (влияет на цвет района на карте)',
        '   • Статус — текущее состояние (В норме / Внимание / Тревога)',
        '   • Район — место события (обязательно)',
        '   • Слой — категория события',
        '   • Описание — подробная информация',
        'Прикрепите файлы (при необходимости):',
        '   • Изображения — фото, скриншоты',
        '   • Документы — Word, PDF и др.',
        'Нажмите «Создать»',
    ]
    for i, step in enumerate(create_steps, 1):
        if step.startswith('   '):
            p = doc.add_paragraph(step.strip())
            p.paragraph_format.left_indent = Cm(1)
        else:
            doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('Редактирование события:', level=3)
    edit_steps = [
        'Найдите событие в списке',
        'Кликните на него для открытия',
        'Нажмите «Редактировать»',
        'Внесите изменения',
        'Нажмите «Сохранить»',
    ]
    for i, step in enumerate(edit_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    important = doc.add_paragraph()
    important.add_run('Важно: ').bold = True
    important.add_run('Редактор может редактировать только созданные им события. Старший редактор и администратор могут редактировать все события.')
    
    doc.add_heading('Удаление события:', level=3)
    delete_steps = [
        'Откройте событие',
        'Нажмите «Удалить»',
        'Подтвердите удаление',
    ]
    for i, step in enumerate(delete_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('Фильтрация событий:', level=3)
    filters = ['По слою', 'По статусу', 'По важности', 'По дате', 'Архивные/активные']
    for f in filters:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('Комментарии:', level=3)
    comments = [
        'В карточке события можно добавлять комментарии',
        'Комментарии видны всем пользователям',
        'Удалить можно только свои комментарии (или любые — для администратора)',
    ]
    for c in comments:
        doc.add_paragraph(c, style='List Bullet')
    
    # 6. Работа со слоями
    doc.add_heading('6. Работа со слоями', level=1)
    doc.add_paragraph('Доступно для: администратор, старший редактор, редактор').italic = True
    
    doc.add_heading('Страница слоёв (/editor/layers)', level=2)
    doc.add_paragraph('Слои используются для категоризации событий.')
    
    doc.add_heading('Структура:', level=3)
    structure = [
        'Слой (верхний уровень)',
        '   └── Подслой (второй уровень)',
        '         └── Под-подслой (третий уровень)',
    ]
    for s in structure:
        doc.add_paragraph(s)
    
    doc.add_heading('Создание слоя:', level=3)
    layer_steps = [
        'Нажмите «Добавить слой»',
        'Введите название',
        'Нажмите «Сохранить»',
    ]
    for i, step in enumerate(layer_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('Создание подслоя:', level=3)
    sublayer_steps = [
        'Найдите родительский слой',
        'Нажмите «Добавить подслой»',
        'Введите название',
        'Нажмите «Сохранить»',
    ]
    for i, step in enumerate(sublayer_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # 7. Отчёты
    doc.add_heading('7. Отчёты', level=1)
    doc.add_paragraph('Доступно для: администратор, старший редактор, редактор').italic = True
    
    doc.add_heading('Страница отчётов (/editor/reports)', level=2)
    
    doc.add_heading('Просмотр событий:', level=3)
    doc.add_paragraph('Таблица со всеми событиями с возможностью фильтрации по различным параметрам.')
    
    doc.add_heading('Экспорт в Excel:', level=3)
    excel_steps = [
        'Настройте фильтры (при необходимости)',
        'Нажмите «Экспорт в Excel»',
        'Файл .xlsx будет скачан',
    ]
    for i, step in enumerate(excel_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('Экспорт в Word:', level=3)
    word_steps = [
        'Выберите конкретное событие',
        'Нажмите «Экспорт в Word»',
        'Файл .docx с детальной информацией будет скачан',
    ]
    for i, step in enumerate(word_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # 8. Администрирование
    doc.add_heading('8. Администрирование', level=1)
    doc.add_paragraph('Доступно только для: администратор').italic = True
    
    doc.add_heading('8.1. Настройки системы (/admin)', level=2)
    settings = [
        'Название управления — отображается на странице обстановки',
        'Выбор региона — определяет карту для отображения',
    ]
    for s in settings:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_heading('8.2. Зоны и устройства (/admin/zones)', level=2)
    doc.add_paragraph('Управление административными зонами:')
    zones = [
        'Создание зон',
        'Привязка районов к зонам',
        'Назначение подразделений',
    ]
    for z in zones:
        doc.add_paragraph(z, style='List Bullet')
    
    doc.add_heading('8.3. Пользователи (/admin/users)', level=2)
    
    doc.add_heading('Создание пользователя:', level=3)
    user_steps = [
        'Нажмите «Добавить пользователя»',
        'Заполните: Логин, Пароль (минимум 6 символов), Полное имя (опционально), Роль',
        'Нажмите «Создать»',
    ]
    for i, step in enumerate(user_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('Управление пользователями:', level=3)
    user_mgmt = [
        'Изменение роли',
        'Сброс пароля',
        'Редактирование данных',
    ]
    for u in user_mgmt:
        doc.add_paragraph(u, style='List Bullet')
    
    doc.add_heading('8.4. Журнал аудита (/admin/journal)', level=2)
    doc.add_paragraph('Просмотр истории действий в системе:')
    audit = [
        'Кто и когда выполнил действие',
        'Тип действия (создание, изменение, удаление)',
        'Затронутые объекты',
    ]
    for a in audit:
        doc.add_paragraph(a, style='List Bullet')
    
    # Навигация
    doc.add_page_break()
    doc.add_heading('Навигация по системе', level=1)
    
    doc.add_heading('Для администратора:', level=2)
    admin_table = doc.add_table(rows=9, cols=2)
    admin_table.style = 'Table Grid'
    admin_nav = [
        ('Кнопка', 'Страница'),
        ('Регион и управление', 'Настройки системы'),
        ('Зоны и устройства', 'Управление зонами'),
        ('Пользователи', 'Управление пользователями'),
        ('Журналирование', 'Журнал аудита'),
        ('Слои', 'Управление слоями'),
        ('События', 'Управление событиями'),
        ('Отчёты', 'Формирование отчётов'),
        ('Обстановка', 'Карта с событиями'),
    ]
    for i, (btn, page) in enumerate(admin_nav):
        row = admin_table.rows[i].cells
        row[0].text = btn
        row[1].text = page
        if i == 0:
            row[0].paragraphs[0].runs[0].font.bold = True
            row[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_heading('Для редакторов:', level=2)
    editor_table = doc.add_table(rows=5, cols=2)
    editor_table.style = 'Table Grid'
    editor_nav = [
        ('Кнопка', 'Страница'),
        ('Слои', 'Управление слоями'),
        ('События', 'Управление событиями'),
        ('Отчёты', 'Формирование отчётов'),
        ('Обстановка', 'Карта с событиями'),
    ]
    for i, (btn, page) in enumerate(editor_nav):
        row = editor_table.rows[i].cells
        row[0].text = btn
        row[1].text = page
        if i == 0:
            row[0].paragraphs[0].runs[0].font.bold = True
            row[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_heading('Для наблюдателей:', level=2)
    doc.add_paragraph('Доступна только страница Обстановка с интерактивной картой.')
    
    # Советы
    doc.add_heading('Горячие клавиши и советы', level=1)
    tips = [
        'Escape — закрытие модальных окон',
        'Клик вне модального окна — закрытие',
        'Используйте фильтры для быстрого поиска нужных событий',
        'События с высокой важностью (7-10) отображаются красным цветом на карте',
    ]
    for t in tips:
        doc.add_paragraph(t, style='List Bullet')
    
    # Поддержка
    doc.add_heading('Поддержка', level=1)
    doc.add_paragraph('При возникновении проблем обратитесь к администратору системы.')
    
    # Футер
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Версия документа: ').bold = True
    footer.add_run('1.0')
    footer2 = doc.add_paragraph()
    footer2.add_run('Дата обновления: ').bold = True
    footer2.add_run('Январь 2026')
    
    # Сохранение
    doc.save('USER_MANUAL.docx')
    print('Документ успешно создан: USER_MANUAL.docx')


if __name__ == '__main__':
    create_manual()
